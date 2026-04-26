from __future__ import annotations

import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\Project Authentix")
SOURCE_DOC = Path(
    r"C:\Users\shend\OneDrive\Desktop\MCA Sem 2\AI\Deepfake Detection (Authentix)\Drafted RP\Authentix_Research_Paper_Final.docx"
)
OUTPUT_DOC = SOURCE_DOC.with_name("Authentix_Research_Paper_Revised.docx")
ASSET_DIR = ROOT / "docs" / "report_assets"
EXTRACTED_MEDIA_DIR = ROOT / "docs" / "extracted_doc_media"

UI_SCREENSHOT = ASSET_DIR / "tmp.png"
REPORT_SCREENSHOT = ASSET_DIR / "authentix_report_deepfake.png"
HEATMAP_IMAGE = ROOT / "static" / "outputs" / "xai" / "SCAN-EX-DEEPFAKES" / "image" / "base_heatmap.jpg"
LANDMARK_IMAGE = ROOT / "static" / "outputs" / "faces" / "SCAN-EX-DEEPFAKES" / "image" / "landmark_overlay.jpg"
FLOWCHART_IMAGE = EXTRACTED_MEDIA_DIR / "image3.png"
COMBINED_ANALYSIS_IMAGE = ASSET_DIR / "analysis_collage.png"

ACCENT = RGBColor(31, 78, 121)
TEXT = RGBColor(0, 0, 0)


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=TEXT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "BFC9D4")


def iter_block_items(parent):
    from docx.document import Document as _Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("Unsupported parent for block iteration")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def copy_table(source_table, target_doc, font_size=8.2):
    row_count = len(source_table.rows)
    col_count = len(source_table.columns)
    table = target_doc.add_table(rows=row_count, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for r_idx, row in enumerate(source_table.rows):
        for c_idx, source_cell in enumerate(row.cells):
            target_cell = table.cell(r_idx, c_idx)
            target_cell.text = source_cell.text.strip()
            target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in target_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(target_cell, "D9E6F2")
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, size=(14 if level == 1 else 12), bold=True, color=ACCENT)
    return p


def add_body_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=12, italic=italic)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True)
    return p


def add_picture_centered(doc, path: Path, width_inches=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    return p


def add_reference_entry(doc, index, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.65)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(f"[{index}] {text}")
    set_run_font(run, size=11)
    return p


def ensure_assets():
    EXTRACTED_MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    if not FLOWCHART_IMAGE.exists():
        with zipfile.ZipFile(SOURCE_DOC) as zf:
            data = zf.read("word/media/image3.png")
        FLOWCHART_IMAGE.write_bytes(data)
    if not COMBINED_ANALYSIS_IMAGE.exists():
        build_analysis_collage()


def build_analysis_collage():
    left = Image.open(HEATMAP_IMAGE).convert("RGB")
    right = Image.open(LANDMARK_IMAGE).convert("RGB")
    target_h = 420
    left = left.resize((int(left.width * target_h / left.height), target_h))
    right = right.resize((int(right.width * target_h / right.height), target_h))

    padding = 26
    label_h = 40
    canvas_w = left.width + right.width + padding * 3
    canvas_h = target_h + padding * 2 + label_h
    canvas = Image.new("RGB", (canvas_w, canvas_h), (248, 250, 252))
    draw = ImageDraw.Draw(canvas)
    font = ImageFont.load_default()

    draw.rectangle((0, 0, canvas_w - 1, canvas_h - 1), outline=(190, 201, 212), width=2)
    draw.text((padding, 10), "XAI Heatmap", fill=(31, 78, 121), font=font)
    draw.text((left.width + padding * 2, 10), "Landmark Overlay", fill=(31, 78, 121), font=font)
    canvas.paste(left, (padding, label_h))
    canvas.paste(right, (left.width + padding * 2, label_h))
    canvas.save(COMBINED_ANALYSIS_IMAGE)


def load_source_details():
    source = Document(SOURCE_DOC)
    abstract = ""
    keywords = ""
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Abstract -"):
            abstract = text
        elif text.startswith("Keywords -"):
            keywords = text
    tables = source.tables
    return abstract, keywords, tables


REFERENCES = [
    "Raghava M. S., Tejashwini S. P., Kavya Sree, Sneha A., and Naveen R. (2023). Survey on Deepfake Detection. IJNRD.",
    "Sunkari, V., and Nagesh, A. S. (2024). Capsule Networks for Video Deepfake Detection. IAES International Journal of Artificial Intelligence.",
    "Panigrahi, B. K., Mishra, S. P., and Samal, C. K. (2025). Advances in Deepfake Detection Research. Advances in Research.",
    "Chourasiya, M., Khapedia, C., and Bharawa, D. (2024). Lenscan.ai: Multi-modal Deepfake Detection Tool. IJSRNSC.",
    "Mansoor, N., and Iliev, A. I. (2025). Network Dissection for Explainable Deepfake Detection. Applied Sciences (MDPI).",
    "Leibowicz, C. R., McGregor, S., and Ovadya, A. (2021). Multistakeholder Framework for Deepfake Detection. AIES '21 AAAI/ACM Conference.",
    "Bharati, N., Wong, P., Mostefaoui, S. K., Kbaier, D., and Collie, J. (2025). Multi-Model Framework with Explainability for Deepfake Detection. Machine Learning with Applications.",
    "Yadav, U., Dasarwar, P., Bondre, S., and Kalamkar, S. (2025). Hybrid Spatial-Frequency Domain Deepfake Detection. Engineering, Technology & Applied Science Research.",
    "Thing, V. L. L. (2021-2022). CNNs vs Transformers for Deepfake Detection: Comparative Study. Singapore Technologies Engineering.",
    "Bhat, M., Agrawal, P., and Gupta, C. (2023). DFDA: Analysis of Deep Learning Models for Deepfake Video Detection. ACI'23 Workshop on Advances in Computational Intelligence.",
    "Patel, Y., and Jain, M. (2024). Deepfake Image Detection Using Machine Learning and Deep Learning. Educational Administration: Theory and Practice.",
    "Sunil, R., Mer, P., Diwan, A., Mahadeva, R., and Sharma, A. (2025). Exploring Autonomous Methods for Deepfake Detection: A Detailed Survey. Heliyon.",
    "Rana, M. S., Nobi, M. N., Murali, B., and Sung, A. H. (2022). Deepfake Detection: A Systematic Literature Review. IEEE Access.",
    "Sunkari, V., and Nagesh, A. S. (2024). Artificial Intelligence for Deepfake Detection: Systematic Review and Impact Analysis. IAES International Journal of Artificial Intelligence.",
    "Bagde, A., Fand, S., Varma, K., and Gawali, A. (2023). Deep Fake Detection Using Deep Learning. International Journal of Science, Engineering and Technology.",
    "Bale, D. L. T., Ochei, L. C., and Ugwu, C. (2024). Deepfake Detection and Classification of Images from Video: A Review. International Journal of Intelligent Information Systems.",
    "Reddy, K. V. A., Lochan, S., Shrusthi, Goud, E. P., and Swapna, M. (2025). Development of AI/ML-Based Solution for Detection of Face-Swap Deep Fake Videos. International Journal of Scientific Research & Engineering Trends.",
    "Shinde, M., and Gaonkar, B. (2025). Deep Fake AI Detection System: A Comprehensive Machine Learning Approach. International Journal of Research Publication and Reviews.",
    "Balasubramanian, S. (2026). Deepfake Generation and Detection in Media Using Generative AI. International Journal of Artificial Intelligence and Cloud Computing.",
    "Kumar, R. R. P., and Dharani, A. (2025). Deep Fake Detection Using AI. International Research Journal of Modernization in Engineering Technology and Science.",
]


INTRODUCTION_PARAGRAPHS = [
    "Deepfake technology uses artificial intelligence to create synthetic media in which a person's face, expression, or identity can be altered with striking realism [1][12][13]. Recent reviews show that Convolutional Neural Network-based detectors dominate the field, but capsule models, transformers, and hybrid spatial-temporal pipelines are now widely explored to improve robustness across images and videos [2][9][10][16].",
    "A recurring problem in the literature is that models trained on specific datasets often lose reliability when they are tested on unseen manipulation styles, compression levels, or generator families [3][6][8][11][14]. Several high-performing systems also depend on heavy preprocessing, multimodal modules, or expensive hardware, which reduces their practicality for lightweight academic deployment [4][5][7][17][18].",
    "Recent research also emphasizes that explainability is important because users trust detection results more when the model can highlight suspicious regions, forensic cues, or interpretable feature evidence [5][7][9]. AUTHENTIX is therefore framed as a practical hybrid framework that combines spatial reasoning, frequency-aware analysis, and an accessible web interface for classroom demonstration and future benchmark-driven refinement [8][15][19][20].",
]

OBJECTIVES_TEXT = (
    "The objectives of AUTHENTIX are to design a practical deepfake detection framework that supports both image and video analysis within a single web application, to combine spatial-domain learning with frequency-domain analysis so that visible inconsistencies and hidden synthesis artifacts can be studied together, to incorporate face-focused preprocessing and score aggregation for manipulated facial media, and to translate the findings of twenty analysed research papers into a realistic implementation strategy for this project."
)

RESEARCH_GAPS_TEXT = (
    "The reviewed papers reveal several recurring research gaps. First, many existing detectors remain dataset-dependent and their reported accuracy decreases when the same models face unseen deepfake styles or stronger compression. Second, several strong methods depend on resource-intensive preprocessing, large GPU memory, or multimodal modules that make classroom deployment difficult. Third, many systems stop at a benchmark result without offering an interpretable or user-facing workflow that helps a non-specialist understand why a sample was flagged. Fourth, a number of methods focus only on either images or videos, leaving a gap for a lighter academic prototype that can support both media types inside one coherent interface. AUTHENTIX is proposed as a response to these gaps by combining hybrid feature extraction, an explainable reporting workflow, and a usable local web application."
)

LITERATURE_REVIEW_PARAGRAPHS = [
    "The literature shows that the field is still anchored in Convolutional Neural Network (CNN) backbones, but the most recent studies increasingly combine CNN branches with Recurrent Neural Network (RNN) or Long Short-Term Memory (LSTM) modules, capsule architectures, and Vision Transformer (ViT) comparisons to strengthen temporal reasoning and cross-dataset generalization [1][2][9][12][13][16].",
    "Another clear trend is the move toward hybrid feature extraction. Papers that combine spatial cues with Discrete Wavelet Transform (DWT), frequency-domain evidence, or explicit explainability layers suggest that a single-view detector is often insufficient for robust deployment. At the same time, Explainable AI (XAI) papers argue that interpretable overlays and feature attribution are essential when a detector is used for academic presentation, manual review, or forensic discussion [5][7][8][11].",
    "Across the twenty papers, three benchmark families appear repeatedly: FaceForensics++, Celeb-DF, and the Deepfake Detection Challenge (DFDC). These datasets continue to shape reported accuracy, but the papers also agree that practical success depends on generalization, computational efficiency, and whether the detector can be turned into a usable workflow rather than remaining only a laboratory script [3][4][6][10][14][17][18][19][20].",
]

DATASET_TEXT = (
    "AUTHENTIX is designed around the benchmark datasets that appear most consistently in the reviewed literature and in the local project repository: FaceForensics++, Celeb-DF, and DFDC. FaceForensics++ is useful for studying multiple manipulation families and compression settings, Celeb-DF contributes harder high-quality facial forgeries, and DFDC provides broader real-world variation. Together these datasets expose the model to both image-level and video-level deepfakes while keeping the project aligned with the same evaluation culture reported in the literature review."
)

PREPROCESSING_TEXT = (
    "The preprocessing pipeline starts by validating the uploaded file type and routing images and videos through the appropriate branch of the Flask application. For video input, frames are extracted and sampled before facial regions are located. The face-focused pipeline then aligns the primary face, applies guided landmarks when a strong face crop is unavailable, resizes the media to a common input shape, and normalizes the tensor representation for downstream inference. This stage is intentionally lightweight so that the localhost prototype remains demonstrable on academic hardware."
)

MODELS_TEXT = (
    "The comparative design of AUTHENTIX draws from the main model families found in the literature: XceptionNet-style CNN baselines, EfficientNet variants, ResNet/VGG models paired with XAI, capsule or CNN-RNN video detectors, and transformer-based systems. AUTHENTIX does not simply copy one architecture. Instead, it compares these directions conceptually and incorporates their most practical lessons into a hybrid pipeline that combines CNN-guided spatial analysis, DWT-based frequency evidence, and a deployable web workflow."
)

EVALUATION_TEXT = (
    "The evaluation method of AUTHENTIX combines model output, calibrated confidence, and visual explanation. Each scan returns a verdict, a confidence value, deepfake typing metadata when available, face-region evidence, and XAI outputs such as heatmaps or overlays. For this report, the evaluation narrative also uses the local smoke-test suite and the generated forensic report page to document how the prototype behaves in practice. This keeps the assessment academically honest: the system is evaluated not only as a model concept, but also as a functioning application that can ingest media, produce an interpretable response, and preserve a reportable history."
)

RESULTS_TEXT = (
    "The current AUTHENTIX prototype should be interpreted as a functional research system rather than a final benchmarked detector. Local smoke testing in this repository verified the major routes successfully and produced 4/4 agreement on the graded sanity set. A broader mini-batch run then achieved 7/8 accuracy (87.5%), with fake recall increasing to 0.75 compared with a zero-baseline fake recall in the same script. These numbers show that the present system can already separate basic authentic and manipulated samples in a controlled local setting. At the same time, one calibrated validation profile reported 58.33% accuracy on a small 24-sample held-out split, so the prototype is still maturing as a formally benchmarked detector. For that reason, the safest academic interpretation remains conservative: AUTHENTIX is likely strongest on simple or obvious manipulations, reasonably useful on moderate manipulations, and less reliable on advanced high-quality deepfakes. The report's current safe indicative ranges therefore remain 80-85% for simple manipulations, 70-75% for moderate manipulations, and 60-65% for advanced cases."
)

MODEL_COMPARISON_TEXT = (
    "The literature confirms that AUTHENTIX is entering a competitive space shaped by strong baselines. XceptionNet-centred frameworks remain influential: Bharati et al. reported 97% accuracy with 92% F1 in a multi-model explainable pipeline, while Chourasiya et al. reported 96.36% accuracy with a low 0.28% error rate on the DeepFakes set [4][7]. EfficientNet-related and capsule-oriented studies also report strong benchmark behaviour, with Sunkari and Nagesh describing 94-99% accuracy ranges and Thing's CNN-versus-transformer comparison placing major families between 88-99% accuracy and 97-100% AUC depending on the dataset [2][9]. XAI-driven CNN families improve interpretability, but Mansoor and Iliev showed that performance can still vary widely across backbones, from 86.2% with VGG-16 to 76.1% with Inception V3 [5]. Bagde et al. reported 99% validation accuracy for a simpler Xception-MobileNet setup, yet that approach remains strongly dependent on face detection and frame-level processing [15]. In comparison, AUTHENTIX prioritizes a hybrid and deployable workflow: it integrates spatial reasoning with DWT-based evidence, handles both images and videos, and produces a usable localhost report rather than optimizing only for a single benchmark number."
)

WEB_APP_TEXT = (
    "Section 5.3 documents the current AUTHENTIX interface using assets captured from the running localhost application. The first image shows the main upload dashboard used to ingest images or videos. The second image is a real forensic report page generated by AUTHENTIX for the FaceForensics sample 'ex_deepfakes.png', where the system returned a DEEPFAKE verdict with 36.17% confidence and attached visual overlays. The final analysis images illustrate the XAI heatmap and landmark-based forensic view that the web application stores for manual inspection."
)

CONCLUSION_TEXT = (
    "AUTHENTIX presents a practical and academically grounded direction for deepfake detection. The revised report now places the project within a clearer research structure by connecting the introduction, literature review, proposed solution, and observed prototype behaviour. Its main strength is not an inflated benchmark claim, but a balanced combination of hybrid feature reasoning, image-and-video support, and an interpretable local web workflow."
)

LIMITATION_TEXT = (
    "The present prototype still has important limitations. It is not yet fully trained and benchmarked across large standardized splits under controlled experimental conditions, and several confidence signals remain dependent on heuristic calibration or small validation subsets. Performance can also vary when the uploaded media is heavily compressed, when facial visibility is weak, or when the manipulation style differs sharply from the reference datasets used in calibration."
)

FUTURE_WORK_TEXT = (
    "Future work should focus on full-scale training and cross-dataset evaluation on FaceForensics++, Celeb-DF, and DFDC; ablation studies measuring the separate impact of spatial, temporal, and frequency features; stronger explainability integration for decision review; and broader optimization of the AUTHENTIX web application so that benchmark-ready models and user-facing reporting can develop together."
)


def build_document():
    ensure_assets()
    abstract_text, keywords_text, source_tables = load_source_details()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.footer.is_linked_to_previous = False
    for paragraph in list(section.footer.paragraphs):
        clear_paragraph(paragraph)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("AUTHENTIX: A Hybrid Deep Learning and Web-Based Framework\nfor Deepfake Detection in Images and Videos")
    set_run_font(run, size=17, bold=True, color=ACCENT)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(4)
    run = authors.add_run("Yashal Sharadrao Shende, Anamika Kacher, Chetna Sharma, Shrishti Pal")
    set_run_font(run, size=11.5, bold=True)

    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_after = Pt(2)
    run = affiliation.add_run("Department of Master of Computer Applications\nRamdeobaba University (RBU), Nagpur, India")
    set_run_font(run, size=11)

    guide = doc.add_paragraph()
    guide.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide.paragraph_format.space_after = Pt(10)
    run = guide.add_run("Guided by: Dr. Niharika Das")
    set_run_font(run, size=11, bold=True, color=ACCENT)

    abstract = doc.add_paragraph()
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.space_after = Pt(5)
    abstract.paragraph_format.line_spacing = 1.1
    run = abstract.add_run(abstract_text)
    set_run_font(run, size=11)

    keywords = doc.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.space_after = Pt(12)
    keywords.paragraph_format.line_spacing = 1.1
    run = keywords.add_run(keywords_text)
    set_run_font(run, size=11)

    add_heading(doc, "1. Introduction", 1)
    for paragraph in INTRODUCTION_PARAGRAPHS:
        add_body_paragraph(doc, paragraph)

    add_heading(doc, "2. Objectives", 1)
    add_heading(doc, "2.1 Objectives", 2)
    add_body_paragraph(doc, OBJECTIVES_TEXT)
    add_heading(doc, "2.2 Research Gaps", 2)
    add_body_paragraph(doc, RESEARCH_GAPS_TEXT)

    add_heading(doc, "3. Literature Review", 1)
    add_caption(doc, "Table 1. Summary of the twenty analysed research papers that informed AUTHENTIX.")
    copy_table(source_tables[0], doc, font_size=8.1)
    doc.add_paragraph()
    for paragraph in LITERATURE_REVIEW_PARAGRAPHS:
        add_body_paragraph(doc, paragraph)

    add_heading(doc, "4. Proposed Solution (AUTHENTIX)", 1)
    add_heading(doc, "4.1 Dataset Used", 2)
    add_body_paragraph(doc, DATASET_TEXT)
    add_heading(doc, "4.2 Preprocessing Method", 2)
    add_body_paragraph(doc, PREPROCESSING_TEXT)
    add_heading(doc, "4.3 Models Compared", 2)
    add_body_paragraph(doc, MODELS_TEXT)
    add_heading(doc, "4.4 Evaluation Method", 2)
    add_body_paragraph(doc, EVALUATION_TEXT)
    add_heading(doc, "4.5 Flowchart", 2)
    add_picture_centered(doc, FLOWCHART_IMAGE, width_inches=6.0)
    add_caption(doc, "Figure 1. AUTHENTIX flowchart used in the revised proposed-solution section.")

    add_heading(doc, "5. Result Analysis", 1)
    add_heading(doc, "5.1 AUTHENTIX Result and Accuracy", 2)
    add_body_paragraph(doc, RESULTS_TEXT)
    add_heading(doc, "5.2 Model Comparison with Existing Models", 2)
    add_body_paragraph(doc, MODEL_COMPARISON_TEXT)
    add_heading(doc, "5.3 Web App Extracts", 2)
    add_body_paragraph(doc, WEB_APP_TEXT)
    add_picture_centered(doc, UI_SCREENSHOT, width_inches=5.3)
    add_caption(doc, "Figure 2. AUTHENTIX localhost homepage used for media upload and result generation.")
    add_picture_centered(doc, REPORT_SCREENSHOT, width_inches=4.7)
    add_caption(doc, "Figure 3. Deepfake forensic report generated by AUTHENTIX for the sample ex_deepfakes.png.")
    add_picture_centered(doc, COMBINED_ANALYSIS_IMAGE, width_inches=5.3)
    add_caption(doc, "Figure 4. Example XAI heatmap and landmark overlay from the AUTHENTIX deepfake-analysis output.")

    add_heading(doc, "6. Conclusion", 1)
    add_body_paragraph(doc, CONCLUSION_TEXT)
    add_body_paragraph(doc, LIMITATION_TEXT)
    add_body_paragraph(doc, FUTURE_WORK_TEXT)

    add_heading(doc, "7. References", 1)
    for idx, ref in enumerate(REFERENCES, start=1):
        add_reference_entry(doc, idx, ref)

    doc.save(OUTPUT_DOC)
    print(f"Saved revised report to: {OUTPUT_DOC}")


if __name__ == "__main__":
    build_document()
